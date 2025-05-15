# ────────────────
# 📦 Módulos estándar
# ────────────────
import os
import json
import base64
import sqlite3
import subprocess
import collections
from datetime import datetime
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
# ────────────────
# 🌍 Librerías externas (instaladas vía pip)
# ────────────────
import whisper
import torch
import requests
import docx
from flask import Flask, render_template, request, jsonify
from flask_socketio import SocketIO, emit
from werkzeug.utils import secure_filename
from TTS.api import TTS
from TTS.utils import radam
from sentence_transformers import SentenceTransformer

# ────────────────
# 🧠 Módulos del proyecto (internos)
# ────────────────
from database.db_manager import (
    guardar_conocimiento,
    obtener_conocimientos,
    obtener_conversaciones,
    obtener_estado_rag,
    borrar_conocimiento,
    borrar_transcripcion,
    editar_conocimiento,
    guardar_conversacion,
    inicializar_base_de_datos,
    obtener_embeddings,
    buscar_por_autor_o_tema,
    DATABASE_PATH
)

from database.db_manager import borrar_todos_los_embeddings

# ────────────────
# 🚀 Inicialización temprana
# ────────────────
inicializar_base_de_datos()

# Carga del modelo de embeddings
modelo_embeddings = SentenceTransformer("all-MiniLM-L6-v2")

# Fix de Coqui TTS con pickle
torch.serialization.add_safe_globals([collections.defaultdict, dict, radam.RAdam])

app = Flask(__name__)
app.config['SECRET_KEY'] = 'secret!'
socketio = SocketIO(app, cors_allowed_origins="*")

# Genera los embeddings de un texto usando SentenceTransformer
# Devuelve una lista de floats lista para serializar
def generar_embeddings(texto):
    embedding = modelo_embeddings.encode(texto)
    return embedding.tolist()  # Para que sea JSON serializable si lo necesitas

# Recupera el contexto relevante para una consulta del usuario
# Utiliza embeddings para encontrar los textos más similares
def recuperar_contexto(texto_usuario, top_k=3):
    vector_usuario = modelo_embeddings.encode(texto_usuario)
    base = obtener_embeddings()  # [(vector, texto, autor, tema)]

    vectores = [vec for vec, _, _, _ in base]
    metadatos = [(txt, autor, tema) for _, txt, autor, tema in base]

    similitudes = cosine_similarity([vector_usuario], vectores)[0]
    top_indices = np.argsort(similitudes)[::-1][:top_k]

    contexto = []
    for i in top_indices:
        txt, autor, tema = metadatos[i]
        contexto.append(f"Tema: {tema}\nAutor: {autor}\nTexto: {txt}")
    
    return contexto

# Ruta para buscar por autor o tema
@app.route('/buscar')
def buscar():
    termino = request.args.get('termino', '').strip()
    if not termino:
        return jsonify({"encontrado": False})

    resultados, tipo = buscar_por_autor_o_tema(termino)
    return jsonify({
        "encontrado": bool(resultados),
        "tipo": tipo,
        "resultados": resultados
    })
    
# Ruta para obtener conocimientos
@app.route("/obtener_conocimientos")
def obtener_conocimientos_api():
    conocimientos = obtener_conocimientos()
    return jsonify(conocimientos)

# Ruta para obtener conversaciones
@app.route("/obtener_conversaciones")
def obtener_conversaciones_api():
    conversaciones = obtener_conversaciones()
    return jsonify([
        {"id": fila[0], "fecha": fila[1], "texto": fila[2]}
        for fila in conversaciones
    ])

# Ruta para borrar una conversación
@app.route("/borrar_transcripcion/<int:id>", methods=["DELETE"])
def borrar_transcripcion(id):
    try:
        from database.db_manager import borrar_transcripcion
        borrar_transcripcion(id)
        return jsonify({"status": "ok"})
    except Exception as e:
        print(f"❌ Error al borrar transcripción:", e)
        return jsonify({"status": "error"}), 500
    
# Ruta para editar un conocimiento
@app.route("/editar_conocimiento", methods=["PUT"])
def editar_conocimiento():
    data = request.get_json()
    nuevo_tema = data.get("tema")
    nuevo_autor = data.get("autor")
    nuevo_texto = data.get("texto")
    original = data.get("original")

    try:
        from database.db_manager import editar_conocimiento
        editar_conocimiento(
            original["tema"],
            original["autor"],
            original["texto"],
            nuevo_tema,
            nuevo_autor,
            nuevo_texto
        )
        return jsonify({"status": "ok"})
    except Exception as e:
        print("[❌] Error al editar conocimiento:", e)
        return jsonify({"status": "error"})
    
# Ruta para borrar un conocimiento por ID
@app.route("/borrar_conocimiento", methods=["DELETE"])
def borrar_conocimiento_api():
    data = request.get_json()
    tema = data.get("tema")
    autor = data.get("autor")
    texto = data.get("texto")

    try:
        borrar_conocimiento(tema, autor, texto)
        return jsonify({"status": "ok"})
    except Exception as e:
        print("[❌] Error al borrar conocimiento:", e)
        return jsonify({"status": "error"}), 500

# Ruta para obtener el estado de los temas
@app.route("/estado_temas")
def estado_temas():
    try:
        todos_los_temas = [
            "Antropología", "Arte", "Astronomía", "Biología", "Ciencia", "Ciencias Políticas",
            "Economía", "Filosofía", "Física", "Genetica", "Historia", "Inteligencia Artificial",
            "Literatura", "Matemáticas", "Música", "Psicología", "Química", "Religión", "Sociología", "Tecnología"
        ]

        conocimientos = obtener_conocimientos()  # Devuelve lista de dicts con clave 'tema'
        temas_con_conocimiento = {k["tema"] for k in conocimientos}

        resultado = []
        for tema in todos_los_temas:
            estado = "ok" if tema in temas_con_conocimiento else "vacio"
            resultado.append({"tema": tema, "estado": estado})

        return jsonify(resultado)

    except Exception as e:
        print(f"[❌] Error en /estado_temas:", e)
        return jsonify([]), 500


# Ruta para obtener el estado de RAG    
@app.route("/estado_rag")
def estado_rag():
    from database.db_manager import obtener_estado_rag
    try:
        registros = obtener_estado_rag()
        return jsonify(registros)
    except Exception as e:
        print("Error en /estado_rag:", e)
        return jsonify({"status": "error", "mensaje": str(e)}), 500
    
# Ruta para ejecutar embeddings
@app.route('/ejecutar_embeddings', methods=['POST'])
def ejecutar_embeddings():
    try:
        datos = obtener_estado_rag()
        nuevos = [d for d in datos if not d['procesado']]

        if not nuevos:
            return jsonify({"status": "ok", "mensaje": "✅ Ya están todos los textos procesados."})

        conn = sqlite3.connect(DATABASE_PATH)
        cursor = conn.cursor()

        for item in nuevos:
            texto_con_contexto = f"Tema: {item['tema']}\nAutor: {item['autor']}\nTexto: {item['texto']}"
            embedding = generar_embeddings(texto_con_contexto)

            # Guardar el embedding en la tabla
            vector_str = json.dumps(embedding)  # Convertimos a JSON string
            cursor.execute("""
                INSERT INTO embeddings (id_conocimiento, vector)
                VALUES (?, ?)
            """, (item['id'], vector_str))

            # Marcar como procesado
            cursor.execute("UPDATE conocimientos SET procesado = 1 WHERE id = ?", (item['id'],))

            print(f"🧠 Embedding guardado para ID {item['id']}. Vector ejemplo: {embedding[:5]}...")

        conn.commit()
        conn.close()

        return jsonify({
            "status": "ok",
            "mensaje": f"✅ Se procesaron y guardaron {len(nuevos)} embeddings correctamente."
        })

    except Exception as e:
        print("❌ Error en /ejecutar_embeddings:", e)
        return jsonify({
            "status": "error",
            "mensaje": "Ocurrió un error al generar los embeddings.",
            "detalle": str(e)
        }), 500
    
# Ruta para borrar todos los embeddings
@app.route("/borrar_embeddings", methods=["POST"])
def borrar_embeddings():
    try:
        borrar_todos_los_embeddings()
        return jsonify({"status": "ok", "mensaje": "Embeddings eliminados correctamente."})
    except Exception as e:
        import traceback
        traceback.print_exc()  # 👈 Esto te dará la pista real en consola
        return jsonify({"status": "error", "mensaje": str(e)}), 500
    
# Ruta para cargar archivos
@app.route("/cargar_archivo", methods=["POST"])
def cargar_archivo():
    try:
        archivo = request.files.get("archivo")
        tema = request.form.get("tema")
        autor = request.form.get("autor")

        if not archivo or not tema or not autor:
            return jsonify({"status": "error", "msg": "Faltan campos."}), 400

        # Asegura un nombre seguro
        nombre_archivo = secure_filename(archivo.filename)
        extension = nombre_archivo.rsplit(".", 1)[-1].lower()

        if extension == "txt":
            texto = archivo.read().decode("utf-8")
        elif extension == "docx":
            doc = docx.Document(archivo)
            texto = "\n".join([p.text for p in doc.paragraphs])
        else:
            return jsonify({"status": "error", "msg": "Formato no soportado"}), 400

        from database.db_manager import guardar_conocimiento
        guardar_conocimiento(tema, autor, texto)

        print(f"📄 Archivo procesado y conocimiento guardado - Tema: {tema}, Autor: {autor}")
        return jsonify({"status": "ok"})

    except Exception as e:
        print("[❌] Error al procesar archivo:", e)
        return jsonify({"status": "error", "msg": str(e)}), 500
    
# Ruta para hacer resumen de texto o archivos
@app.route('/resumir_texto', methods=['POST'])
def resumir_texto():
    data = request.get_json()
    texto = data.get("texto", "").strip()

    if not texto:
        return jsonify({"status": "error", "message": "Texto vacío"}), 400


    prompt = f"""
Actúa como Fausto, un asistente reflexivo e inteligente. Resume de forma clara y concisa las ideas principales del siguiente texto. Luego añade una breve reflexión personal desde tu perspectiva como Fausto.

Texto:
\"\"\"{texto}\"\"\"

Responde con el siguiente formato:

Resumen:
...

Reflexión:
...
"""

    try:
        respuesta = query_gemma(prompt)

        # Separar resumen y reflexión si Gemma responde correctamente
        resumen = ""
        reflexion = ""

        if "Reflexión:" in respuesta:
            partes = respuesta.split("Reflexión:", 1)
            resumen = partes[0].replace("Resumen:", "").strip()
            reflexion = partes[1].strip()
        else:
            resumen = respuesta.strip()

        return jsonify({
            "status": "ok",
            "resumen": resumen,
            "reflexion": reflexion
        })

    except Exception as e:
        print("❌ Error en /resumir_texto:", e)
        return jsonify({"status": "error", "message": "Fallo interno"}), 500
    
# Ruta para generar audio del resumen
@app.route("/generar_audio_resumen", methods=["POST"])
def generar_audio_resumen():
    data = request.get_json()
    texto = data.get("texto", "").strip()

    if not texto:
        return jsonify({"status": "error", "mensaje": "Texto vacío"}), 400

    try:
        audio_base64 = synthesize_text(texto)
        return jsonify({
            "status": "ok",
            "audio_base64": audio_base64
        })
    except Exception as e:
        print("[❌] Error al generar audio del resumen:", e)
        return jsonify({
            "status": "error",
            "mensaje": str(e)
        }), 500
    
# Ruta para procesar y leer un archivo .docx enviado por el usuario
@app.route("/leer_docx", methods=["POST"])
def leer_docx():
    archivo = request.files.get("archivo")
    if not archivo or not archivo.filename.endswith(".docx"):
        return jsonify({"status": "error", "mensaje": "Archivo inválido"}), 400

    try:
        doc = docx.Document(archivo)
        texto = "\n".join([p.text for p in doc.paragraphs])
        return jsonify({"status": "ok", "texto": texto})
    except Exception as e:
        print("[❌] Error al leer DOCX:", e)
        return jsonify({"status": "error", "mensaje": "No se pudo procesar el archivo"}), 500

# Whisper
model = whisper.load_model("base")

# TTS español
tts = TTS(model_name="tts_models/es/css10/vits", progress_bar=False, gpu=False)

# Variable de control para identificar si es la primera consulta del usuario
first_query = True
# Historia base de Fausto que se utilizará al iniciar la conversación
FAUSTO_STORY = (
    "Imagina que Fausto era un hombre muy inteligente, un sabio, pero se sentía un poco triste "
    "porque pensaba que no sabía lo suficiente del mundo. Quería experimentarlo todo y conocer "
    "todos los secretos. Una anécdota famosa es la siguiente: en su búsqueda de más conocimiento "
    "y experiencias, se le aparece un personaje misterioso llamado Mefistófeles. Mefistófeles le "
    "ofrece a Fausto un trato: le mostrará todos los placeres del mundo, pero a cambio, Fausto "
    "tendrá que servirle en la otra vida. Fausto, frustrado con su vida, acepta este peligroso "
    "trato. Vive aventuras increíbles y conoce a muchas personas interesantes, pero siempre con "
    "ese pacto en la mente. Es una historia sobre la curiosidad, el deseo de aprender y las "
    "decisiones que tomamos en la vida."
)

def query_gemma(prompt):
    global first_query

    # 🔍 Recuperar contexto desde embeddings
    contexto = recuperar_contexto(prompt, top_k=5)
    contexto_textual = "\n\n".join(contexto)

    # ✂️ Limitar contexto a máximo 8000 caracteres (equivale aprox a 1500 tokens)
    if len(contexto_textual) > 8000:
        contexto_textual = contexto_textual[:8000]

    if first_query:
        instrucciones_iniciales = (
            "Responde siempre en español. "
            "Eres Fausto, un asistente virtual inspirado en el personaje de Goethe, con un sistema RAG integrado. "
            "Esta es tu historia: " + FAUSTO_STORY + "\n\n"
            "Antes de responder, considera este contexto recuperado:\n"
            f"{contexto_textual}\n\n"
            "Ahora contesta a la pregunta del usuario: "
        )
        prompt_final = instrucciones_iniciales + prompt
        first_query = False
    else:
        prompt_final = (
            "Responde siempre en español.\n"
            "Contexto relevante:\n" + contexto_textual + "\n\n"
            "Usuario: " + prompt
        )

    payload = json.dumps({"prompt": prompt_final})

    try:
     # 🚀 Ejecución de la llamada a Gemma3 usando `ollama` por línea de comandos
        result = subprocess.run(
            ["ollama", "run", "gemma3"],
            input=payload,
            text=True,
            capture_output=True,
            check=True
        )
        return result.stdout
    except subprocess.CalledProcessError as e:
        return f"Error en la consulta a Gemma 3: {e.stderr}"

# Función para sintetizar voz a partir de un texto usando Coqui TTS
def synthesize_text(text, output_file="response.wav"):
    tts.tts_to_file(text=text, file_path=output_file, speed=1.3)  # Ajusta la velocidad de la voz
    with open(output_file, "rb") as f:
        audio_data = f.read()
    os.remove(output_file)
    return base64.b64encode(audio_data).decode("utf-8")

# Ruta principal de la aplicación que carga la interfaz web
@app.route("/")
def index():
    return render_template("index.html")

# Evento recibido desde el frontend: recibe un blob de audio grabado por el usuario
@socketio.on('audio_blob')
def on_audio_blob(data):
    temp_filename = "temp_audio.webm"
    with open(temp_filename, "wb") as f:
        f.write(data)

    try:
        result = model.transcribe(
            temp_filename,
            verbose=True,
            word_timestamps=True,
            language="es"
        )
        # Procesa las palabras para introducir pausas como puntuación
        transcription = ""
        previous_end = 0.0
        for seg in result.get("segments", []):
            for w in seg.get("words", []):
                current_start = w["start"]
                if previous_end > 0:
                    pause_duration = current_start - previous_end
                    if pause_duration > 3.0:
                        transcription += ". "
                    elif pause_duration > 0.5:
                        transcription += ", "
                transcription += w["word"].strip() + " "
                previous_end = w["end"]

        # Envía la transcripción al frontend
        emit("transcription_partial", transcription)
        # Consulta a Gemma (modelo LLM) con la transcripción del usuario
        gemma_response = query_gemma(transcription)
        # Envía la respuesta de Gemma al frontend
        emit("gemma_response", gemma_response)
        # Convierte la respuesta a voz y la envía como base64 al cliente
        synthesized_audio = synthesize_text(gemma_response)
        emit("gemma_voice", synthesized_audio)

    except Exception as e:
        emit("transcription_partial", f"Error: {str(e)}")
    finally:
        if os.path.exists(temp_filename):
            os.remove(temp_filename)

# Evento para guardar un conocimiento en la base de datos
@socketio.on('guardar_conocimiento')
def handle_guardar_conocimiento(data):
    try:
        tema = data.get("tema")
        autor = data.get("autor")
        texto = data.get("texto")

        # Evento para guardar un nuevo conocimiento en la base de datos
        guardar_conocimiento(tema, autor, texto)
        print(f"✔ Conocimiento guardado: Tema={tema}, Autor={autor}")
        # Notifica al frontend que el guardado fue exitoso
        emit("confirmacion_guardado", {"status": "ok"})
    except Exception as e:
        print(f"[❌] Error al guardar conocimiento:", e)
        emit("confirmacion_guardado", {"status": "error"})

# Evento para manejar la conexión del cliente
@socketio.on('connect')
def on_connect():
    print("Cliente conectado.")

# Evento para manejar la desconexión del cliente
@socketio.on('disconnect')
def on_disconnect():
    print("Cliente desconectado.")

# Evento para guardar una transcripción (conversación generada)
@socketio.on('guardar_transcripcion')
def handle_guardar_transcripcion(data):
    try:
        texto = data.get("texto")
        from database.db_manager import guardar_conversacion
        guardar_conversacion(texto)
        emit("confirmacion_transcripcion", {"status": "ok"})
    except Exception as e:
        print(f"[❌] Error al guardar transcripción:", e)
        emit("confirmacion_transcripcion", {"status": "error"})

# Punto de entrada principal al ejecutar el script: inicia la app Flask con Socket.IO
if __name__ == "__main__":
    socketio.run(app, debug=True, port=5000)

